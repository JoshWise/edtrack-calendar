"""
EdTrack Calendar Module - FastAPI Application

This module provides REST API endpoints for scraping curriculum content,
processing calendar data, and integrating with the main EdTrack application.
"""

from fastapi import FastAPI, HTTPException, BackgroundTasks, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import List, Dict, Any, Optional
import pandas as pd
import asyncio
from datetime import datetime
import logging
import io
import os
from pathlib import Path

from calendar_scraper import EdTrackCalendarScraper
from calendar_processor import EdTrackCalendarProcessor
# Note: Calendar module is stateless - it doesn't interact with database directly
# It just processes data and returns JSON for the main EdTrack app to use

# Helper function to convert numpy types to Python types
def convert_numpy_types(obj):
    """Convert numpy types to Python native types for JSON serialization"""
    if isinstance(obj, dict):
        return {k: convert_numpy_types(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [convert_numpy_types(item) for item in obj]
    elif hasattr(obj, 'item'):  # numpy scalar
        return obj.item()
    elif hasattr(obj, 'tolist'):  # numpy array
        return obj.tolist()
    else:
        return obj

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(
    title="EdTrack Calendar Module",
    description="API for scraping and processing curriculum content and school calendars",
    version="1.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure this for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic models for API requests/responses
class ScrapeCalendarRequest(BaseModel):
    calendar_url: str = Field(..., description="URL to scrape calendar from")
    school_id: int = Field(..., description="School ID from EdTrack database")

class ScrapeLessonsRequest(BaseModel):
    lesson_url: str = Field(..., description="URL to scrape lesson content from")
    class_id: int = Field(..., description="Class ID from EdTrack database")

class ScrapeAndScheduleRequest(BaseModel):
    lesson_url: str = Field(..., description="URL to scrape lesson content from")
    calendar_url: str = Field(..., description="URL to scrape calendar from")
    class_id: int = Field(..., description="Class ID from EdTrack database")
    school_id: int = Field(..., description="School ID from EdTrack database")
    hours_per_day: int = Field(default=1, description="Number of class hours per day")

class ImportDataRequest(BaseModel):
    calendar_data: List[Dict[str, Any]] = Field(..., description="Calendar data to import")
    lesson_data: List[Dict[str, Any]] = Field(..., description="Lesson data to import")
    target_data: List[Dict[str, Any]] = Field(..., description="Learning target data to import")
    school_id: int = Field(..., description="School ID from EdTrack database")
    class_id: int = Field(..., description="Class ID from EdTrack database")

class APIResponse(BaseModel):
    status: str = Field(..., description="Response status")
    message: str = Field(..., description="Response message")
    data: Optional[Dict[str, Any]] = Field(None, description="Response data")
    summary: Optional[Dict[str, Any]] = Field(None, description="Summary statistics")

# Health check endpoint
@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

# Inspect file structure endpoint
@app.post("/inspect-file", response_model=APIResponse)
async def inspect_file(
    file: UploadFile = File(...)
):
    """
    Inspect the structure of an uploaded file to help with processing
    
    Args:
        file: File to inspect
        
    Returns:
        APIResponse with file structure information
    """
    try:
        logger.info(f"Inspecting file structure: {file.filename}")
        
        # Read file content
        file_content = await file.read()
        file_extension = Path(file.filename).suffix.lower()
        
        # Process based on file type
        if file_extension in ['.csv', '.txt']:
            df = pd.read_csv(io.BytesIO(file_content))
        elif file_extension in ['.xlsx', '.xls']:
            df = pd.read_excel(io.BytesIO(file_content))
        elif file_extension == '.json':
            df = pd.read_json(io.BytesIO(file_content))
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")
        
        # Inspect file structure
        processor = EdTrackCalendarProcessor()
        structure = processor.inspect_file_structure(df)
        
        return APIResponse(
            status="success",
            message=f"File structure inspected successfully",
            data={"structure": structure}
        )
        
    except Exception as e:
        logger.error(f"Error inspecting file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to inspect file: {str(e)}")

# Scrape school calendar from URL
@app.post("/scrape-calendar", response_model=APIResponse)
async def scrape_calendar(request: ScrapeCalendarRequest):
    """
    Scrape school calendar from a URL
    
    Args:
        request: Calendar scraping request
        
    Returns:
        APIResponse with calendar data
    """
    try:
        logger.info(f"Scraping calendar from: {request.calendar_url}")
        
        async with EdTrackCalendarScraper() as scraper:
            calendar_df = await scraper.scrape_school_calendar(request.calendar_url, request.school_id)
        
        if calendar_df.empty:
            raise HTTPException(status_code=404, detail="No calendar data found")
        
        # Process calendar data
        processor = EdTrackCalendarProcessor()
        processed_calendar = processor.process_calendar_data(calendar_df, request.school_id)
        
        # Analyze calendar
        analysis = processor.analyze_calendar(processed_calendar)
        
        return APIResponse(
            status="success",
            message="Calendar scraped successfully",
            data={
                "calendar": processed_calendar.to_dict('records')
            },
            summary=analysis
        )
        
    except Exception as e:
        logger.error(f"Error scraping calendar: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to scrape calendar: {str(e)}")

# Upload calendar file (CSV, Excel, etc.)
@app.post("/upload-calendar", response_model=APIResponse)
async def upload_calendar(
    file: UploadFile = File(...),
    school_id: int = Form(...),
    use_visual_parser: str = Form("false")
):
    """
    Upload and process a calendar file (CSV, Excel, PDF, DOCX)
    
    Args:
        file: Calendar file to upload
        school_id: School ID from EdTrack database
        
    Returns:
        APIResponse with calendar data
    """
    try:
        use_visual = use_visual_parser.lower() == 'true'
        logger.info(f"Processing calendar file: {file.filename}, school_id: {school_id}, visual_parser: {use_visual}")
        
        # Read file content first
        file_content = await file.read()
        file_extension = Path(file.filename).suffix.lower()
        
        # Check if visual parser is requested and available
        if use_visual and file_extension in ['.pdf', '.png', '.jpg', '.jpeg']:
            try:
                from calendar_visual_parser import VisualCalendarParser, check_visual_parser_available
                
                available, msg = check_visual_parser_available()
                if not available:
                    logger.warning(f"Visual parser not available: {msg}. Falling back to text extraction.")
                    use_visual = False
                else:
                    # Save file temporarily for visual parser
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                        tmp_file.write(file_content)
                        tmp_path = tmp_file.name
                    
                    # Use visual parser
                    vparser = VisualCalendarParser()
                    df = vparser.parse_pdf_calendar(tmp_path, school_id)
                    processed_calendar = vparser.to_edtrack_format(df)
                    
                    # Clean up temp file
                    os.unlink(tmp_path)
                    
                    # Analyze calendar
                    processor = EdTrackCalendarProcessor()
                    analysis = processor.analyze_calendar(processed_calendar)
                    
                    return APIResponse(
                        status="success",
                        message=f"Calendar file '{file.filename}' processed with VISUAL parser (colors, symbols detected)",
                        data={
                            "calendar": processed_calendar.to_dict('records')
                        },
                        summary=convert_numpy_types(analysis)
                    )
            except ImportError as e:
                logger.warning(f"Visual parser dependencies not installed: {e}. Using text extraction.")
                use_visual = False
            except Exception as e:
                logger.error(f"Visual parser failed: {e}. Falling back to text extraction.")
                use_visual = False
        
        # Standard text-based parsing (default or fallback)
        # Note: file_content already read above
        
        # Process based on file type
        if file_extension in ['.csv', '.txt']:
            # Read CSV file
            calendar_df = pd.read_csv(io.BytesIO(file_content))
        elif file_extension in ['.xlsx', '.xls']:
            # Read Excel file
            calendar_df = pd.read_excel(io.BytesIO(file_content))
        elif file_extension == '.json':
            # Read JSON file
            calendar_df = pd.read_json(io.BytesIO(file_content))
        elif file_extension == '.rtf':
            # Read RTF file - extract text and parse
            import re
            rtf_text = file_content.decode('utf-8', errors='ignore')
            # Simple RTF stripping
            text = re.sub(r'\\[a-z]+\d*\s?', ' ', rtf_text)
            text = re.sub(r'[{}]', '', text)
            text = re.sub(r'\\\'[0-9a-fA-F]{2}', '', text)
            text = re.sub(r'\s+', ' ', text)
            
            # Try to parse as CSV-like data from the extracted text
            from io import StringIO
            try:
                calendar_df = pd.read_csv(StringIO(text))
            except:
                # If CSV parsing fails, try to extract structured data
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                if len(lines) > 0:
                    # Create basic calendar structure from lines
                    calendar_df = pd.DataFrame({
                        'date': [],
                        'day_type': [],
                        'description': []
                    })
                else:
                    raise HTTPException(status_code=400, detail="Could not extract calendar data from RTF file")
        elif file_extension in ['.docx', '.doc']:
            # Read DOCX file using python-docx
            from docx import Document
            from io import BytesIO
            doc = Document(BytesIO(file_content))
            
            # Extract text from paragraphs and tables
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
            
            # Extract data from tables if present
            table_data = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Skip empty rows
                        table_data.append(row_data)
            
            # Try to create DataFrame from table data
            if table_data and len(table_data) > 1:
                # Assume first row is header
                calendar_df = pd.DataFrame(table_data[1:], columns=table_data[0])
            else:
                # Try to parse text content as CSV
                from io import StringIO
                full_text = '\n'.join(text_content)
                try:
                    calendar_df = pd.read_csv(StringIO(full_text))
                except:
                    raise HTTPException(status_code=400, detail="Could not extract calendar data from DOCX file. Ensure the file contains a table or CSV-formatted text.")
        elif file_extension == '.pdf':
            # Read PDF file using PyPDF2
            from PyPDF2 import PdfReader
            from io import BytesIO
            pdf_reader = PdfReader(BytesIO(file_content))
            
            # Extract text from all pages
            text_content = []
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
            
            full_text = '\n'.join(text_content)
            
            # Try to parse as CSV
            from io import StringIO
            try:
                calendar_df = pd.read_csv(StringIO(full_text))
            except:
                raise HTTPException(status_code=400, detail="Could not extract calendar data from PDF file. Consider converting to CSV or Excel format.")
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")
        
        if calendar_df.empty:
            raise HTTPException(status_code=404, detail="No calendar data found in file")
        
        # Add school_id if not present
        if 'school_id' not in calendar_df.columns:
            calendar_df['school_id'] = school_id
        
        # Process calendar data
        processor = EdTrackCalendarProcessor()
        processed_calendar = processor.process_calendar_data(calendar_df, school_id)
        
        # Analyze calendar
        analysis = processor.analyze_calendar(processed_calendar)
        
        return APIResponse(
            status="success",
            message=f"Calendar file '{file.filename}' processed successfully",
            data={
                "calendar": processed_calendar.to_dict('records')
            },
            summary=convert_numpy_types(analysis)
        )
        
    except Exception as e:
        logger.error(f"Error processing calendar file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to process calendar file: {str(e)}")

# Scrape lesson content from URL
@app.post("/scrape-lessons", response_model=APIResponse)
async def scrape_lessons(request: ScrapeLessonsRequest):
    """
    Scrape lesson content from a URL
    
    Args:
        request: Lesson scraping request
        
    Returns:
        APIResponse with lesson data
    """
    try:
        logger.info(f"Scraping lessons from: {request.lesson_url}")
        
        async with EdTrackCalendarScraper() as scraper:
            lessons_df = await scraper.scrape_lesson_content(request.lesson_url, request.class_id)
        
        if lessons_df.empty:
            raise HTTPException(status_code=404, detail="No lesson data found")
        
        # Extract learning targets
        processor = EdTrackCalendarProcessor()
        targets_df = processor.create_learning_targets_from_lessons(lessons_df)
        
        return APIResponse(
            status="success",
            message="Lessons scraped successfully",
            data={
                "lessons": lessons_df.to_dict('records'),
                "targets": targets_df.to_dict('records')
            },
            summary={
                "total_lessons": len(lessons_df),
                "total_targets": len(targets_df),
                "duration_types": lessons_df['duration_type'].value_counts().to_dict() if 'duration_type' in lessons_df.columns else {}
            }
        )
        
    except Exception as e:
        logger.error(f"Error scraping lessons: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to scrape lessons: {str(e)}")

# Upload lesson file (CSV, Excel, PDF, DOCX)
@app.post("/upload-lessons", response_model=APIResponse)
async def upload_lessons(
    file: UploadFile = File(...),
    class_id: int = Form(...)
):
    """
    Upload and process a lesson file (CSV, Excel, PDF, DOCX)
    
    Args:
        file: Lesson file to upload
        class_id: Class ID from EdTrack database
        
    Returns:
        APIResponse with lesson data
    """
    try:
        logger.info(f"Processing lesson file: {file.filename}")
        
        # Read file content
        file_content = await file.read()
        file_extension = Path(file.filename).suffix.lower()
        
        # Process based on file type
        if file_extension in ['.csv', '.txt']:
            # Read CSV file
            lessons_df = pd.read_csv(io.BytesIO(file_content))
        elif file_extension in ['.xlsx', '.xls']:
            # Read Excel file
            lessons_df = pd.read_excel(io.BytesIO(file_content))
        elif file_extension == '.json':
            # Read JSON file
            lessons_df = pd.read_json(io.BytesIO(file_content))
        elif file_extension == '.rtf':
            # Read RTF file - extract lesson structure
            import re
            rtf_text = file_content.decode('utf-8', errors='ignore')
            # Simple RTF stripping
            text = re.sub(r'\\[a-z]+\d*\s?', ' ', rtf_text)
            text = re.sub(r'[{}]', '', text)
            text = re.sub(r'\\\'[0-9a-fA-F]{2}', '', text)
            text = re.sub(r'\s+', ' ', text)
            
            # Extract lesson-like patterns (numbers, titles, etc.)
            lesson_patterns = [
                r'(\d+\.\d+\.?\d*)[:\s]+([^\n]+)',
                r'Lesson (\d+)[:\s]+([^\n]+)',
                r'Activity (\d+\.\d+\.?\d*)[:\s]+([^\n]+)',
                r'LESSON (\d+)[:\s]+([^\n]+)',
                r'UNIT (\d+)[:\s]+([^\n]+)',
            ]
            
            lessons_found = []
            for pattern in lesson_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    lesson_num = match.group(1).strip()
                    title = match.group(2).strip()
                    if title and len(title) > 3 and len(title) < 200:
                        lessons_found.append({
                            'lesson_number': lesson_num,
                            'title': title,
                            'class_id': class_id
                        })
            
            if lessons_found:
                lessons_df = pd.DataFrame(lessons_found)
                # Remove duplicates
                lessons_df = lessons_df.drop_duplicates(subset=['lesson_number', 'title'])
            else:
                # Fallback: try CSV parsing
                from io import StringIO
                try:
                    lessons_df = pd.read_csv(StringIO(text))
                except:
                    raise HTTPException(status_code=400, detail="Could not extract lesson data from RTF file. Try using CSV or Excel format.")
        elif file_extension in ['.docx', '.doc']:
            # Read DOCX file using python-docx
            from docx import Document
            from io import BytesIO
            doc = Document(BytesIO(file_content))
            
            # Extract text from paragraphs
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
            
            # Extract data from tables if present
            table_data = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        table_data.append(row_data)
            
            # Try to create DataFrame from table data first
            if table_data and len(table_data) > 1:
                # Assume first row might be header
                try:
                    lessons_df = pd.DataFrame(table_data[1:], columns=table_data[0])
                except:
                    lessons_df = pd.DataFrame(table_data)
            else:
                # Extract lessons from text using patterns
                import re
                full_text = '\n'.join(text_content)
                
                lesson_patterns = [
                    r'(\d+\.\d+\.?\d*)[:\s]+([^\n]+)',
                    r'Lesson (\d+)[:\s]+([^\n]+)',
                    r'Activity (\d+\.\d+\.?\d*)[:\s]+([^\n]+)',
                    r'LESSON (\d+)[:\s]+([^\n]+)',
                    r'UNIT (\d+)[:\s]+([^\n]+)',
                ]
                
                lessons_found = []
                for pattern in lesson_patterns:
                    matches = re.finditer(pattern, full_text, re.IGNORECASE)
                    for match in matches:
                        lesson_num = match.group(1).strip()
                        title = match.group(2).strip()
                        if title and len(title) > 3 and len(title) < 200:
                            lessons_found.append({
                                'lesson_number': lesson_num,
                                'title': title,
                                'class_id': class_id
                            })
                
                if lessons_found:
                    lessons_df = pd.DataFrame(lessons_found)
                    lessons_df = lessons_df.drop_duplicates(subset=['lesson_number', 'title'])
                else:
                    raise HTTPException(status_code=400, detail="Could not extract lesson data from DOCX file. Ensure the file contains structured lesson information.")
        elif file_extension == '.pdf':
            # Read PDF file using PyPDF2
            from PyPDF2 import PdfReader
            from io import BytesIO
            import re
            
            pdf_reader = PdfReader(BytesIO(file_content))
            
            # Extract text from all pages
            text_content = []
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
            
            full_text = '\n'.join(text_content)
            
            # Extract lessons using patterns
            lesson_patterns = [
                r'(\d+\.\d+\.?\d*)[:\s]+([^\n]+)',
                r'Lesson (\d+)[:\s]+([^\n]+)',
                r'Activity (\d+\.\d+\.?\d*)[:\s]+([^\n]+)',
                r'LESSON (\d+)[:\s]+([^\n]+)',
                r'UNIT (\d+)[:\s]+([^\n]+)',
            ]
            
            lessons_found = []
            for pattern in lesson_patterns:
                matches = re.finditer(pattern, full_text, re.IGNORECASE)
                for match in matches:
                    lesson_num = match.group(1).strip()
                    title = match.group(2).strip()
                    if title and len(title) > 3 and len(title) < 200:
                        lessons_found.append({
                            'lesson_number': lesson_num,
                            'title': title,
                            'class_id': class_id
                        })
            
            if lessons_found:
                lessons_df = pd.DataFrame(lessons_found)
                lessons_df = lessons_df.drop_duplicates(subset=['lesson_number', 'title'])
            else:
                raise HTTPException(status_code=400, detail="Could not extract lesson data from PDF file. Consider converting to CSV or Excel format.")
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")
        
        if lessons_df.empty:
            raise HTTPException(status_code=404, detail="No lesson data found in file")
        
        # Add class_id if not present
        if 'class_id' not in lessons_df.columns:
            lessons_df['class_id'] = class_id
        
        # Ensure required columns exist
        if 'lesson_number' not in lessons_df.columns:
            lessons_df['lesson_number'] = range(1, len(lessons_df) + 1)
        if 'title' not in lessons_df.columns:
            lessons_df['title'] = lessons_df.apply(lambda row: f"Lesson {row['lesson_number']}", axis=1)
        if 'status' not in lessons_df.columns:
            lessons_df['status'] = 'planned'
        
        # Extract learning targets
        processor = EdTrackCalendarProcessor()
        targets_df = processor.create_learning_targets_from_lessons(lessons_df)
        
        return APIResponse(
            status="success",
            message=f"Lesson file '{file.filename}' processed successfully",
            data={
                "lessons": lessons_df.to_dict('records'),
                "targets": targets_df.to_dict('records')
            },
            summary={
                "total_lessons": len(lessons_df),
                "total_targets": len(targets_df),
                "file_type": file_extension
            }
        )
        
    except Exception as e:
        logger.error(f"Error processing lesson file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to process lesson file: {str(e)}")

# Scrape and schedule lessons with calendar
@app.post("/scrape-and-schedule", response_model=APIResponse)
async def scrape_and_schedule(request: ScrapeAndScheduleRequest):
    """
    Scrape lessons and calendar, then schedule lessons across school days
    
    Args:
        request: Combined scraping and scheduling request
        
    Returns:
        APIResponse with scheduled lessons and calendar data
    """
    try:
        logger.info(f"Scraping and scheduling from lesson: {request.lesson_url} and calendar: {request.calendar_url}")
        
        async with EdTrackCalendarScraper() as scraper:
            # Scrape both lesson and calendar data concurrently
            lessons_task = scraper.scrape_lesson_content(request.lesson_url, request.class_id)
            calendar_task = scraper.scrape_school_calendar(request.calendar_url, request.school_id)
            
            lessons_df, calendar_df = await asyncio.gather(lessons_task, calendar_task)
        
        if lessons_df.empty:
            raise HTTPException(status_code=404, detail="No lesson data found")
        if calendar_df.empty:
            raise HTTPException(status_code=404, detail="No calendar data found")
        
        # Process data
        processor = EdTrackCalendarProcessor()
        
        # Process calendar
        processed_calendar = processor.process_calendar_data(calendar_df, request.school_id)
        
        # Schedule lessons
        scheduled_lessons = processor.process_lessons_for_scheduling(
            lessons_df, processed_calendar, request.hours_per_day, request.class_id
        )
        
        # Extract learning targets
        learning_targets = processor.create_learning_targets_from_lessons(lessons_df)
        
        # Create lesson-target mappings
        lesson_target_mappings = processor.create_lesson_target_mappings(scheduled_lessons, learning_targets)
        
        # Validate schedule
        validation = processor.validate_schedule(scheduled_lessons, processed_calendar)
        
        # Analyze calendar
        calendar_analysis = processor.analyze_calendar(processed_calendar)
        
        return APIResponse(
            status="success",
            message="Lessons scraped and scheduled successfully",
            data={
                "calendar": processed_calendar.to_dict('records'),
                "lessons": scheduled_lessons.to_dict('records'),
                "targets": learning_targets.to_dict('records'),
                "mappings": lesson_target_mappings.to_dict('records'),
                "validation": validation
            },
            summary={
                "total_lessons": len(scheduled_lessons),
                "total_targets": len(learning_targets),
                "school_days": calendar_analysis.get('school_days', 0),
                "total_calendar_days": calendar_analysis.get('total_days', 0),
                "schedule_valid": validation.get('valid', False),
                "warnings": validation.get('warnings', []),
                "errors": validation.get('errors', [])
            }
        )
        
    except Exception as e:
        logger.error(f"Error scraping and scheduling: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to scrape and schedule: {str(e)}")

# Get curriculum metadata
@app.post("/curriculum-metadata", response_model=APIResponse)
async def get_curriculum_metadata(request: ScrapeLessonsRequest):
    """
    Get metadata about curriculum content without scraping full lessons
    
    Args:
        request: Lesson scraping request (for URL)
        
    Returns:
        APIResponse with curriculum metadata
    """
    try:
        logger.info(f"Getting curriculum metadata from: {request.lesson_url}")
        
        async with EdTrackCalendarScraper() as scraper:
            metadata = await scraper.scrape_curriculum_metadata(request.lesson_url)
        
        return APIResponse(
            status="success",
            message="Curriculum metadata retrieved successfully",
            data={"metadata": metadata}
        )
        
    except Exception as e:
        logger.error(f"Error getting curriculum metadata: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get metadata: {str(e)}")

# Process existing data (for testing/debugging)
@app.post("/process-data", response_model=APIResponse)
async def process_existing_data(request: ImportDataRequest):
    """
    Process existing lesson and calendar data
    
    Args:
        request: Data processing request
        
    Returns:
        APIResponse with processed data
    """
    try:
        logger.info("Processing existing data")
        
        # Convert data to DataFrames
        lessons_df = pd.DataFrame(request.lesson_data)
        calendar_df = pd.DataFrame(request.calendar_data)
        
        if lessons_df.empty or calendar_df.empty:
            raise HTTPException(status_code=400, detail="No data provided")
        
        # Process data
        processor = EdTrackCalendarProcessor()
        
        # Process calendar
        processed_calendar = processor.process_calendar_data(calendar_df, request.school_id)
        
        # Schedule lessons
        scheduled_lessons = processor.process_lessons_for_scheduling(
            lessons_df, processed_calendar, hours_per_day=1, class_id=request.class_id
        )
        
        # Extract learning targets
        learning_targets = processor.create_learning_targets_from_lessons(lessons_df)
        
        # Create lesson-target mappings
        lesson_target_mappings = processor.create_lesson_target_mappings(scheduled_lessons, learning_targets)
        
        # Validate schedule
        validation = processor.validate_schedule(scheduled_lessons, processed_calendar)
        
        return APIResponse(
            status="success",
            message="Data processed successfully",
            data={
                "calendar": processed_calendar.to_dict('records'),
                "lessons": scheduled_lessons.to_dict('records'),
                "targets": learning_targets.to_dict('records'),
                "mappings": lesson_target_mappings.to_dict('records'),
                "validation": validation
            },
            summary={
                "total_lessons": len(scheduled_lessons),
                "total_targets": len(learning_targets),
                "schedule_valid": validation.get('valid', False)
            }
        )
        
    except Exception as e:
        logger.error(f"Error processing data: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to process data: {str(e)}")

# Get API documentation
@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "EdTrack Calendar Module API",
        "version": "1.0.0",
        "docs": "/docs",
        "health": "/health",
        "endpoints": {
            "inspect_file": "POST /inspect-file",
            "scrape_calendar": "POST /scrape-calendar",
            "scrape_lessons": "POST /scrape-lessons",
            "scrape_and_schedule": "POST /scrape-and-schedule",
            "curriculum_metadata": "POST /curriculum-metadata",
            "process_data": "POST /process-data",
            "upload_calendar": "POST /upload-calendar",
            "upload_lessons": "POST /upload-lessons"
        }
    }

# Error handlers
@app.exception_handler(404)
async def not_found_handler(request, exc):
    from fastapi.responses import JSONResponse
    return JSONResponse(
        status_code=404,
        content={"status": "error", "message": "Endpoint not found"}
    )

@app.exception_handler(500)
async def internal_error_handler(request, exc):
    from fastapi.responses import JSONResponse
    return JSONResponse(
        status_code=500,
        content={"status": "error", "message": "Internal server error"}
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
